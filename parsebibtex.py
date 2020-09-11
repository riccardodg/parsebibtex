#!/usr/bin/python3

import bibtexparser
from bibtexparser.bwriter import BibTexWriter
from bibtexparser.bibdatabase import BibDatabase

import docx
import re
from docx import Document


import argparse
import sys
from datetime import datetime
import os


def argparser(routine):
    # optional input list variables

    example_text_0 = """Usage:

    python3 {} bib_file.bib R""".format(
        routine
    )

    example_text_1 = """Attention:  if NO -n flag is provided, -r and -ne are ignored: usage example:

    python3 {} -p SIR -c US
    python3 {} -p SIR -c Italy  -r Toscana, Piemonte
    python3 {} -p SIR -c Italy  -ne Lombardia, Veneto""".format(
        routine, routine, routine
    )
    usage_text = example_text_0  # + "\n\n" + example_text_1
    parser = argparse.ArgumentParser(
        prog=routine,
        description="Manages bibtex file and formats its content according to different output formats.",
        epilog=usage_text,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    # mandatory arguments
    parser.add_argument(
        "bib_file",
        action="store",
        # dest='plot_type',
        help="Provide the bib file",
        metavar="BIB_FILE",
        type=str,
        default="",
    )
    parser.add_argument(
        "type",
        action="store",
        # dest='plot_type',
        help="Provide a type. R for ricercatore, T for tecnologo",
        metavar="TYPE",
        choices=["R", "T"],
        type=str,
        default="",
    )
    
    parser.add_argument(
        "name",
        action="store",
        # dest='plot_type',
        help="Provide the name",
        metavar="NAME",
        type=str,
        default="",
    )

    args = parser.parse_args()
    # get the values

    bib_file = args.bib_file
    type = args.type
    name=args.name

    return bib_file, type,name
'''
article in journal
'''
def parse_article(n, bib_item):
    ret_text=""
    wrong_item=None
    #table structure
    new_keys_journal=['Nr ', 'Tipologia prodotto ','Elenco autori ','Titolo ','Rivista ','Codice identificativo (ISSN) ','anno pubblicazione ',
    'Indice di classificazione ', 'Impact Factor rivista ','ruolo svolto ', 'numero citazioni ','Altre informazioni ']
    journal_table={}
    #id
    ret_text=new_keys_journal[0]+str(n)
    journal_table[0]=ret_text
    
    #Tipologia prodotto
    ret_text=new_keys_journal[1]+"Articolo in Rivista"
    journal_table[1]=ret_text
    
    #Elenco autori
    ret_text=new_keys_journal[2]+bib_item['author']
    journal_table[2]=ret_text
    
    #Titolo
    ret_text=new_keys_journal[3]+bib_item['title']
    regex = r"\\'"
    test_str = ret_text
    subst = "'"
    ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
    journal_table[3]=ret_text
    
    #Rivista
    ret_text=new_keys_journal[4]+bib_item['journal']
    journal_table[4]=ret_text
    
    #ISSN
    try:
        ret_text=new_keys_journal[5]+bib_item['issn']
    except:
        wrong_item=bib_item
        print("WI ",bib_item['title'])
    journal_table[5]=ret_text
    
    #year
    ret_text=new_keys_journal[6]+bib_item['year']
    journal_table[6]=ret_text
    
    #fields from 7 to 10
    ret_text=new_keys_journal[7]
    journal_table[7]=ret_text
    
    ret_text=new_keys_journal[8]
    journal_table[8]=ret_text
    
    ret_text=new_keys_journal[9]
    journal_table[9]=ret_text
    
    ret_text=new_keys_journal[10]
    journal_table[10]=ret_text
    
    #managing altre informazioni
    # I put url, doi, abstract
    ret_text=new_keys_journal[11]+" "
    doi=""
    url=""
    abs=""
    try:
        doi="doi: "+bib_item['doi']
    except:
        print("item, ",n, "no doi")

    try:
        url="url: "+bib_item['url']
    except:
            print("item, ", n, "no url")
    try:
        abs="abstract: "+bib_item['abstract']
    except:
        print("item, ",n, "no abstract")
    ret_text=ret_text+' '.join([doi,url,abs])
    journal_table[11]=ret_text
    return journal_table, wrong_item
    
    
'''
article in proceedings with or w/o ISBN
different tables with ISBN wrt w/o ISBN
'''
def parse_inproceedings(n, bib_item):
    isbn=""
    ret_text=""
    proceedings_tables={}
    new_keys_isbn=['Nr ', 'Tipologia prodotto ','Elenco autori ','Titolo ','Codice identificativo (ISBN) ','anno pubblicazione ','Altre informazioni ']
    
    new_keys_no_isbn=['Nr ', 'Tipologia prodotto ','Titolo ','Descrizione ','Elenco autori ','Ruolo svolto ','anno pubblicazione ','Altre informazioni ']
    
    try:
        isbn=bib_item['isbn']
    except:
        isbn=""
    
    if not isbn == "":
        #with isbn
        #id
        ret_text=new_keys_isbn[0]+str(n)
        proceedings_tables[0]=ret_text
        
        #Tipologia prodotto
        ret_text=new_keys_isbn[1]+"Contributo in Conferenza o Convegno"
        proceedings_tables[1]=ret_text
        
        #Elenco autori
        ret_text=new_keys_isbn[2]+bib_item['author']
        proceedings_tables[2]=ret_text
        
        #Titolo
        ret_text=new_keys_isbn[3]+bib_item['title']
        regex = r"\\'"
        test_str = ret_text
        subst = "'"
        ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
        proceedings_tables[3]=ret_text
        
        #ISBN
        ret_text=new_keys_isbn[4]+bib_item['isbn']
        proceedings_tables[4]=ret_text
        
        #Year
        ret_text=new_keys_isbn[5]+bib_item['year']
        proceedings_tables[5]=ret_text
        
        #Altre info
        # I put url, booktitle, abstract
        
        booktitle=""
        url=""
        abs=""
        try:
            booktitle=bib_item['booktitle']
        except:
            print("item, ",n, "no booktitle")

        try:
            url="url: "+bib_item['url']
        except:
                print("item, ", n, "no url")
        try:
            abs="abstract: "+bib_item['abstract']
            regex = r"\\'"
            test_str = abs
            subst = "'"
            ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
                      
            regex = r"\\\""
            test_str = ret_text
            subst = "\""
            ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
        except:
            print("item, ",n, "no abstract")
        
        ret_text=new_keys_isbn[6]+" "
        ret_text=ret_text+' '.join([booktitle,url,abs])
        proceedings_tables[6]=ret_text
        
    else:
        #w/o isbn
        #id
        ret_text=new_keys_no_isbn[0]+str(n)
        proceedings_tables[0]=ret_text
        
        #Tipologia prodotto
        ret_text=new_keys_no_isbn[1]+"Contributo in Conferenza o Convegno (Senza ISBN)"
        proceedings_tables[1]=ret_text
        
        #Titolo
        ret_text=new_keys_no_isbn[2]+bib_item['title']
        regex = r"\\'"
        test_str = ret_text
        subst = "'"
        ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
        proceedings_tables[2]=ret_text
        
        #Descrizione
        ret_text=new_keys_no_isbn[3]
        proceedings_tables[3]=ret_text
        
        #Elenco autori
        ret_text=new_keys_no_isbn[4]+bib_item['author']
        proceedings_tables[4]=ret_text
        
        #Ruolo svolto
        ret_text=new_keys_no_isbn[5]
        proceedings_tables[5]=ret_text
        
        #Year
        ret_text=new_keys_no_isbn[6]+bib_item['year']
        proceedings_tables[6]=ret_text
        
        #Altre info
        booktitle=""
        url=""
        abs=""
        try:
            booktitle=bib_item['booktitle']
        except:
            print("item, ",n, "no booktitle")

        try:
            url="url: "+bib_item['url']
        except:
                print("item, ", n, "no url")
        try:
            abs="abstract: "+bib_item['abstract']
            regex = r"\\'"
            test_str = abs
            subst = "'"
            ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
            
            regex = r"\\\""
            test_str = ret_text
            subst = "\""
            ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
        except:
            print("item, ",n, "no abstract")

        ret_text=ret_text+' '.join([booktitle,url,abs])
        ret_text=new_keys_no_isbn[7]+" "+ret_text
        proceedings_tables[7]=ret_text
        
        
    return proceedings_tables

'''
article in book with or w/o ISSN
'''
def parse_inbooks(n, bib_item):
    ret_text=""
    books_tables={}
    new_keys_book=['Nr ', 'Tipologia prodotto ','Elenco autori ','Titolo ','Codice identificativo (ISBN) ','anno pubblicazione ','Altre informazioni ']
    


    
    #id
    ret_text=new_keys_book[0]+str(n)
    books_tables[0]=ret_text
        
    #Tipologia prodotto
    ret_text=new_keys_book[1]+"Contributo in Collana"
    books_tables[1]=ret_text
    
    #Elenco autori
    ret_text=new_keys_book[2]+bib_item['author']
    books_tables[2]=ret_text
    
    #Titolo
    ret_text=new_keys_book[3]+bib_item['title']
    regex = r"\\'"
    test_str = ret_text
    subst = "'"
    ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
    books_tables[3]=ret_text
    
    #ISBN
    ret_text=new_keys_book[4]+bib_item['isbn']
    books_tables[4]=ret_text
        
    #Year
    ret_text=new_keys_book[5]+bib_item['year']
    books_tables[5]=ret_text
    
    #Altre info
    # I put url, booktitle, abstract, doi issn, note
    
    booktitle=""
    url=""
    abs=""
    issn=""
    doi=""
    note=""
    try:
        booktitle=bib_item['booktitle']
    except:
        print("item, ",n, "no booktitle")

    try:
        url="url: "+bib_item['url']
    except:
        print("item, ", n, "no url")
    
    try:
        doi="doi: "+bib_item['doi']
    except:
        print("item, ",n, "no doi")
    
    try:
        issn="ISSN: "+bib_item['issn']
    except:
        print("item, ", n, "no ISSN")
    
    try:
        issn="Note: "+bib_item['note']
    except:
        print("item, ", n, "no NOTE")
    
    try:
        abs="abstract: "+bib_item['abstract']
        regex = r"\\'"
        test_str = abs
        subst = "'"
        ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
                    
        regex = r"\\\""
        test_str = ret_text
        subst = "\""
        ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
    except:
        print("item, ",n, "no abstract")

    ret_text=new_keys_book[6]+" "
    ret_text=ret_text+' '.join([booktitle,url,abs,doi,issn,note])
    books_tables[6]=ret_text
        
    
        
    return books_tables

'''
techreport
'''
def parse_techreport(n, bib_item):
    
    ret_text=""
    techreport_tables={}
    new_keys_techreport=['Nr ', 'Tipologia prodotto ','Titolo ','Descrizione ','Elenco autori ','Ruolo svolto ','anno pubblicazione ','Altre informazioni ']
    
    #w/o isbn
    #id
    ret_text=new_keys_techreport[0]+str(n)
    techreport_tables[0]=ret_text
    
    #Tipologia prodotto
    ret_text=new_keys_techreport[1]+"Deliverable di Progetto "
    techreport_tables[1]=ret_text
    
    #Titolo
    ret_text=new_keys_techreport[2]+bib_item['title']
    regex = r"\\'"
    test_str = ret_text
    subst = "'"
    ret_text = re.sub(regex, subst, test_str, 0, re.MULTILINE)
    techreport_tables[2]=ret_text
    
    #Descrizione
    ret_text=new_keys_techreport[3]
    techreport_tables[3]=ret_text
        
    #Elenco autori
    ret_text=new_keys_techreport[4]+bib_item['author']
    techreport_tables[4]=ret_text
    
    #Ruolo svolto
    ret_text=new_keys_techreport[5]
    techreport_tables[5]=ret_text
    
    #Year
    ret_text=new_keys_techreport[6]+bib_item['year']
    techreport_tables[6]=ret_text
    
    #Altre info
    ret_text=new_keys_techreport[7]
    techreport_tables[7]=ret_text
    
    return techreport_tables
    
'''
parse the bib file
'''
def parse_bibtext(bib_file, type):
    tables_dict={}
    wrong_items={}
    temp={}
    if type=="R":
       tables_dict,wrong_items=parse_bibtext_R(bib_file)
    if type=="T":
       tables_dict,wrong_items=parse_bibtext_T(bib_file)
    
    print("#items ",len(tables_dict))
    print("#wrong items ",len(wrong_items))
    return tables_dict, wrong_items

'''
parse the bib file 4 R
'''
def parse_bibtext_R(bib_file):
    tables_dict={}
    wrong_items={}
    temp={}
    with open(bib_file) as bibtex_file:
        bib_database = bibtexparser.load(bibtex_file)

    n = 1
    for bib_item in bib_database.entries:
        print(bib_item['ENTRYTYPE'], n)
        
        if bib_item['ENTRYTYPE']=="inproceedings":
            print("call inproceedings ",n)
            temp=parse_inproceedings(n,bib_item)
            tables_dict[n]=temp
            #wrong_items[n]=wi
        if bib_item['ENTRYTYPE']=="article":
            print("call article ",n, bib_item['title'])
            temp,wi=parse_article(n,bib_item)
            tables_dict[n]=temp
            if wi is not None:
                wrong_items[n]=wi
        if bib_item['ENTRYTYPE']=="techreport":
            print("call techreport ",n)
            temp=parse_techreport(n,bib_item)
            tables_dict[n]=temp
        if bib_item['ENTRYTYPE']=="inbook":
            print("call inbook ",n)
            temp=parse_inbooks(n,bib_item)
            tables_dict[n]=temp
        if bib_item['ENTRYTYPE']=="book":
            print("call inbook (book) ",n)
            temp=parse_inbooks(n,bib_item)
            tables_dict[n]=temp
        if bib_item['ENTRYTYPE']=="misc":
            print("call inproceedings (misc) ",n)
            temp=parse_inproceedings(n,bib_item)
            tables_dict[n]=temp
            
        n = n + 1
    print("#items ",n, len(bib_database.entries))
    print("#wrong items ",n, len(wrong_items))
    return tables_dict, wrong_items


'''
parse the bib file 4 T
'''
def parse_bibtext_T(bib_file):
    tables_dict={}
    wrong_items={}
    temp={}
    with open(bib_file) as bibtex_file:
        bib_database = bibtexparser.load(bibtex_file)

    n = 1
    for bib_item in bib_database.entries:
        print(bib_item['ENTRYTYPE'], n)
        
        if bib_item['ENTRYTYPE']=="inproceedings":
            print("call inproceedings ",n)
            temp=parse_inproceedings(n,bib_item)
            tables_dict[n]=temp
            #wrong_items[n]=wi
        if bib_item['ENTRYTYPE']=="article":
            print("call article ",n, bib_item['title'])
            temp,wi=parse_article(n,bib_item)
            tables_dict[n]=temp
            if wi is not None:
                wrong_items[n]=wi
        if bib_item['ENTRYTYPE']=="techreport":
            print("call techreport ",n)
            temp=parse_techreport(n,bib_item)
            tables_dict[n]=temp
        if bib_item['ENTRYTYPE']=="inbook":
            print("call inbook ",n)
            temp=parse_inbooks(n,bib_item)
            tables_dict[n]=temp
            
        n = n + 1
    print("#items ",n, len(bib_database.entries))
    print("#wrong items ",n, len(wrong_items))
    return tables_dict, wrong_items



'''
print doc with tables
'''
def print_doc(doc, doc_name,dict):
    for key in dict.keys():
        val=dict[key]
        num_rows=len(val)
        num_cols=1
        table = doc.add_table(num_rows, num_cols,style='TableGrid')
        #print(dict[key], len(dict[key]))
        for r in range(num_rows):
            cell = table.cell(r, num_cols-1)
            c_text=val[r]
            cell.text=c_text
            
        doc.add_paragraph('')
    doc.save(doc_name)
    

'''
main
'''
def main():
    new_document = Document()
    routine = sys.argv[0]
    #output file
    new_doc_name='./bib/bibfile'
    suffix='.docx'
    
    (bib_file, type,name) = argparser(routine)
    print(bib_file, type,name)
    new_doc_name=new_doc_name+"_"+type+"_"+name+suffix
    tables_dict,y=parse_bibtext(bib_file, type)
    ##print(y)
    print_doc(new_document,new_doc_name,tables_dict)

main()
'''
ds = [NOISBN, JOURNAL]
d = {}
for k in NOISBN.keys():
	try:	
		d[k] = tuple(d[k] for d in ds)
	except:
		print("pass NOISBN ",k)	 
for k in JOURNAL.keys():
	try:	
		d[k] = tuple(d[k] for d in ds)
	except:
		print("pass JOURNAL ",k)
'''
